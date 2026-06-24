from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docs" / "transportation_a2a_adk_code_explanation.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B7C9DA"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_fixed_table_width(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor.from_string("555555")
    subtitle.paragraph_format.space_after = Pt(12)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Transportation Data A2A ADK Agent").bold = True
    sub = doc.add_paragraph(style="Subtitle")
    sub.add_run("Code Explanation and Function Reference")
    meta = doc.add_paragraph()
    meta.add_run("Generated for: ").bold = True
    meta.add_run("A2A ADK agent connecting Gemini Enterprise, BigQuery, and Google Cloud Storage")


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_fixed_table_width(table, [9360])
    set_table_borders(table, "D9E3ED")
    cell = table.cell(0, 0)
    set_cell_fill(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(f" {body}")


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths=(2200, 7160)) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_fixed_table_width(table, list(widths))
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Explanation"
    for c in hdr:
        set_cell_fill(c, LIGHT_BLUE)
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def file_section(doc: Document, path: str, purpose: str, details: list[str]) -> None:
    doc.add_heading(path, level=2)
    add_callout(doc, "Purpose:", purpose)
    add_bullets(doc, details)


def function_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    widths = [2600, 2900, 3860]
    set_fixed_table_width(table, widths)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, name in enumerate(["Function/Class", "Responsibility", "Important Behavior"]):
        hdr[i].text = name
        set_cell_fill(hdr[i], LIGHT_BLUE)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for name, responsibility, behavior in rows:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = responsibility
        cells[2].text = behavior
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)


def build_doc() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This project is a Python Google ADK agent exposed through the Agent-to-Agent "
        "(A2A) protocol. It is designed for Gemini Enterprise to answer questions "
        "using an allowlist of transportation BigQuery tables and documents stored "
        "in Google Cloud Storage."
    )
    add_callout(
        doc,
        "Current scope:",
        "The agent can list configured sources, inspect BigQuery schemas, preview and query allowed tables, list GCS objects, read text files, extract text from PDFs, and search PDF manuals for page-level evidence.",
    )
    add_kv_table(
        doc,
        [
            ("Cloud platform", "Google Cloud Run hosts the A2A service container."),
            ("Agent framework", "Google ADK creates the Gemini-powered agent and tool interface."),
            ("A2A protocol", "A2A exposes an Agent Card and JSON-RPC-compatible interaction endpoint for Gemini Enterprise."),
            ("BigQuery", "Configured tables: bridge, crash, eilis, road, and traffic in the bridge_inventory dataset."),
            ("Cloud Storage", "Configured bucket: YOUR_BUCKET_NAME. Used for text and PDF document retrieval."),
            ("PDF support", "pypdf extracts text from PDFs so the agent can answer from bridge manuals and cite pages."),
        ],
    )

    doc.add_heading("2. Runtime Flow", level=1)
    add_numbered(
        doc,
        [
            "Gemini Enterprise discovers the A2A Agent Card from the Cloud Run URL.",
            "A user asks a question in Gemini Enterprise.",
            "The A2A request enters the Starlette app in data_a2a_agent.agent.",
            "The custom request converter removes unsupported A2A metadata before ADK calls Gemini, avoiding the Enterprise part_metadata error.",
            "The ADK Agent chooses tools from data_a2a_agent.tools based on the prompt.",
            "Tools call BigQuery or Cloud Storage, apply guardrails, and return structured data.",
            "The agent writes a natural-language answer with source context such as SQL, object name, or PDF page.",
        ],
    )

    doc.add_heading("3. File-by-File Reference", level=1)
    file_section(
        doc,
        "data_a2a_agent/__init__.py",
        "Marks data_a2a_agent as a Python package.",
        [
            "Contains only a package docstring.",
            "Allows imports such as data_a2a_agent.agent and data_a2a_agent.tools.",
        ],
    )

    file_section(
        doc,
        "data_a2a_agent/config.py",
        "Centralizes environment variable parsing and validated settings.",
        [
            "Defines table metadata structures used by the BigQuery tools.",
            "Reads Cloud Run and local .env-style variables through os.getenv.",
            "Validates BigQuery table identifiers and aliases before tools use them.",
            "Adds PDF extraction limits: max PDF bytes, max PDF pages, and max returned PDF text characters.",
        ],
    )
    function_table(
        doc,
        [
            ("_TABLE_RE", "Regex constant", "Accepts project.dataset.table or dataset.table. It rejects malformed table ids before queries are built."),
            ("BigQueryTable", "Frozen dataclass for one allowed table", "Stores alias, project, dataset, and table; immutability prevents accidental runtime mutation."),
            ("BigQueryTable.full_id", "Builds canonical table id", "Returns project.dataset.table for BigQuery API calls and validation."),
            ("BigQueryTable.sql_ref", "Builds SQL-safe table reference", "Wraps full_id in backticks so generated SELECT statements can reference the table safely."),
            ("Settings", "Frozen dataclass for all runtime config", "Holds GCP project, location, model, BigQuery allowlist, GCS bucket/prefix, limits, and port."),
            ("Settings.bigquery_table_map", "Alias lookup helper", "Returns a dict of alias to BigQueryTable so tools can resolve bridge, crash, road, etc."),
            ("load_settings()", "Main config loader", "Reads environment variables, applies defaults, parses tables, and returns a Settings object."),
            ("_parse_bigquery_tables()", "Table allowlist parser", "Splits comma-separated table ids and aliases; validates format, default project, duplicate aliases, and alias syntax."),
            ("_parse_bool()", "Boolean parser", "Treats 1, true, yes, y, and on as True. Everything else is False."),
            ("_parse_int()", "Positive integer parser", "Reads numeric limits and raises ValueError if a value is zero or negative."),
        ],
    )

    file_section(
        doc,
        "data_a2a_agent/tools.py",
        "Provides all callable tools that the ADK agent can use.",
        [
            "Implements BigQuery read-only operations with allowlist validation.",
            "Implements Cloud Storage object listing and text/PDF reading.",
            "Adds SQL guardrails, GCS prefix guardrails, and PDF size/page/text limits.",
            "Uses lazy imports for google.cloud and pypdf so module import remains lightweight.",
        ],
    )
    function_table(
        doc,
        [
            ("_BLOCKED_SQL", "Mutation/DDL keyword regex", "Blocks obvious ALTER, DELETE, DROP, INSERT, UPDATE, MERGE, EXPORT, and permission statements."),
            ("_FROM_OR_JOIN", "Table reference regex", "Finds fully qualified tables used after FROM or JOIN to enforce the allowlist."),
            ("list_configured_sources()", "Source inventory tool", "Returns configured BigQuery aliases/full ids and the GCS bucket/prefix."),
            ("describe_bigquery_tables()", "Schema inspection tool", "Calls BigQuery get_table for selected aliases and returns schema fields, descriptions, and row counts."),
            ("preview_bigquery_table()", "Small table preview tool", "Resolves an alias, builds SELECT * LIMIT n, and delegates to run_bigquery_select()."),
            ("run_bigquery_select()", "Read-only query tool", "Validates SQL, enforces a LIMIT, sets maximum_bytes_billed, runs BigQuery, and returns rows."),
            ("list_gcs_objects()", "GCS listing tool", "Lists objects under the configured bucket and prefix with name, size, update time, and content type."),
            ("read_gcs_text_object()", "Text reader tool", "Downloads a byte range from a GCS object and decodes it as UTF-8 with replacement for invalid bytes."),
            ("read_gcs_pdf_object()", "PDF page extraction tool", "Downloads a PDF, extracts selected pages using pypdf, applies page and character limits, and returns page text."),
            ("search_gcs_pdf_object()", "PDF search tool", "Extracts PDF text page by page, searches query terms, and returns page snippets with context."),
            ("_select_tables()", "Internal table selection", "Returns all configured tables or resolves requested aliases."),
            ("_get_table()", "Alias resolver", "Raises a clear error listing known aliases when an unknown alias is requested."),
            ("_validate_select_sql()", "SQL guardrail", "Allows only SELECT/WITH, blocks semicolons and mutation keywords, and rejects non-allowlisted tables."),
            ("_ensure_limit()", "Result-size guardrail", "Wraps SQL in SELECT * FROM (...) LIMIT n if no trailing LIMIT exists."),
            ("_bounded_limit()", "Limit clamp", "Ensures requested limits stay between 1 and configured maximum."),
            ("_combine_prefix()", "GCS prefix builder", "Combines base and requested prefixes while preserving the allowed prefix boundary."),
            ("_assert_blob_allowed()", "GCS path guardrail", "Rejects '..' path segments and objects outside the configured prefix."),
            ("_require_bucket()", "GCS config validator", "Raises a clear error when GCS_BUCKET_NAME is missing."),
            ("_assert_pdf_name()", "PDF extension check", "Requires object names to end with .pdf before PDF extraction is attempted."),
            ("_download_pdf_reader()", "PDF download/parser", "Checks GCS object size, downloads bytes, and creates a pypdf PdfReader from memory."),
            ("_extract_pdf_page_text()", "PDF page text extractor", "Extracts text for one page and normalizes missing extraction to an empty string."),
            ("_bigquery_client()", "BigQuery client factory", "Creates a BigQuery client for the configured billing/project context."),
            ("_import_bigquery()", "Lazy BigQuery import", "Imports google.cloud.bigquery only when a BigQuery tool is called."),
            ("_import_storage()", "Lazy Storage import", "Imports google.cloud.storage only when a GCS tool is called."),
        ],
    )

    file_section(
        doc,
        "data_a2a_agent/agent.py",
        "Builds the ADK agent, A2A server, Agent Card, and health endpoint.",
        [
            "Defines the agent instructions shown to Gemini, including table/domain rules and PDF behavior.",
            "Registers tools from tools.py with the ADK Agent.",
            "Creates a custom A2A executor so incoming request metadata is stripped before Gemini Enterprise receives it.",
            "Builds the Agent Card using both transport and protocol_binding compatibility paths.",
            "Exposes /healthz for Cloud Run startup and operational checks.",
        ],
    )
    function_table(
        doc,
        [
            ("settings", "Module-level Settings object", "Loaded once at import so agent instructions and default A2A URL can reflect environment configuration."),
            ("_build_instruction()", "Agent system instruction builder", "Creates detailed runtime instructions for BigQuery, GCS, PDF search, schema inspection, and source citation behavior."),
            ("root_agent", "ADK Agent instance", "Uses settings.agent_model, transportation instructions, and the full tool list."),
            ("_get_user_id()", "A2A user resolver", "Uses authenticated call-context user if present; otherwise builds a stable A2A_USER_<context_id> value."),
            ("_convert_request_without_part_metadata()", "Gemini Enterprise compatibility converter", "Converts A2A message parts to GenAI Content and sets RunConfig(custom_metadata=None) to avoid part_metadata errors."),
            ("_build_a2a_app()", "Starlette/A2A app factory", "Creates Runner, InMemory services, A2aAgentExecutor, DefaultRequestHandler, A2AStarletteApplication, and adds A2A routes immediately."),
            ("_agent_card()", "Agent Card factory", "Advertises agent name, description, tags, examples, capabilities, and endpoint URL. Tries current and legacy A2A interface field names."),
            ("a2a_app", "ASGI app exported to uvicorn", "This is the object used by Docker CMD: data_a2a_agent.agent:a2a_app."),
            ("healthz()", "Cloud Run health route", "Returns status ok, agent name, configured BigQuery table count, and whether a GCS bucket is configured."),
        ],
    )

    doc.add_heading("4. Deployment and Configuration Files", level=1)
    add_kv_table(
        doc,
        [
            ("Dockerfile", "Builds a Python 3.12 slim image, installs requirements, copies data_a2a_agent, exposes 8080, and starts uvicorn."),
            ("requirements.txt", "Declares google-adk[a2a], BigQuery/Storage clients, uvicorn, Starlette, sse-starlette, and pypdf."),
            (".env.example", "Template for local/Cloud Run environment variables, including BigQuery tables, GCS bucket, and PDF limits."),
            (".gcloudignore", "Prevents .env, .venv, __pycache__, pyc files, and local tooling artifacts from uploading to Cloud Run build context."),
            (".gitignore", "Prevents secrets, virtualenvs, caches, build outputs, and OS metadata from being committed."),
            (".vscode/tasks.json", "Defines VS Code tasks for Windows dependency install, local run, tests, and Cloud Run deploy."),
            ("README.md", "General project overview, local run steps, deploy steps, IAM guidance, and sample prompts."),
            ("WINDOWS.md", "Windows/VS Code-specific setup and deploy instructions."),
            ("data_a2a_agent/Archive.zip", "A project archive artifact. It is not imported by the runtime code."),
        ],
    )

    doc.add_heading("5. Script Reference", level=1)
    add_kv_table(
        doc,
        [
            ("scripts/set-env.ps1.txt", "Loads key=value pairs from .env into the current PowerShell process environment, skipping comments and blanks."),
            ("scripts/run-local.ps1.txt", "Creates .venv with Python 3.12 if needed, installs dependencies, loads .env, sets PORT, and runs uvicorn locally."),
            ("scripts/deploy-cloud-run.ps1.txt", "Loads .env, validates required variables, writes a temporary Cloud Run env YAML, and deploys with gcloud run deploy."),
            ("scripts/deploy-cloud-run.sh.txt", "Bash version of the deploy helper; writes a temp env-vars file and deploys to Cloud Run."),
        ],
    )

    doc.add_heading("6. Test Reference", level=1)
    file_section(
        doc,
        "tests/test_config.py",
        "Unit tests for configuration parsing.",
        [
            "test_parse_bigquery_tables_with_aliases checks table ids and aliases resolve to expected full ids.",
            "test_parse_bigquery_tables_rejects_alias_length_mismatch ensures alias count must match table count.",
            "test_parse_bigquery_tables_requires_project ensures dataset.table requires GOOGLE_CLOUD_PROJECT.",
        ],
    )
    file_section(
        doc,
        "tests/test_tools.py",
        "Unit tests for SQL and GCS guardrails.",
        [
            "test_validate_select_sql_accepts_allowed_table verifies allowlisted SELECT queries pass.",
            "test_validate_select_sql_rejects_mutation verifies non-SELECT mutation SQL is rejected.",
            "test_validate_select_sql_rejects_unconfigured_table verifies queries cannot reference unconfigured tables.",
            "test_assert_blob_allowed_limits_prefix verifies GCS object access stays within the configured prefix.",
        ],
    )

    doc.add_heading("7. Key Operational Guardrails", level=1)
    add_kv_table(
        doc,
        [
            ("BigQuery allowlist", "Only configured full table ids from BIGQUERY_TABLES may be queried."),
            ("Read-only SQL", "The code permits SELECT/WITH only and blocks common mutation/DDL keywords."),
            ("Cost cap", "BIGQUERY_MAX_BYTES_BILLED limits BigQuery scan cost per query job."),
            ("Result cap", "BIGQUERY_DEFAULT_LIMIT and BIGQUERY_MAX_LIMIT control returned row counts."),
            ("GCS prefix", "GCS_PREFIX limits object access inside the bucket."),
            ("PDF size cap", "GCS_MAX_PDF_BYTES prevents very large PDF downloads."),
            ("PDF page cap", "GCS_MAX_PDF_PAGES limits extraction/search work per tool call."),
            ("PDF text cap", "GCS_MAX_PDF_TEXT_CHARS limits how much extracted text is returned to Gemini."),
            ("IAM boundary", "Runtime IAM roles still provide the hard security boundary; application checks are defense in depth."),
        ],
    )

    doc.add_heading("8. Common User Questions the Agent Should Handle", level=1)
    add_bullets(
        doc,
        [
            "What data sources do you have access to?",
            "Describe the schema for the bridge and crash tables.",
            "How many bridge records are available by county or route?",
            "Show crash counts by severity and month for the latest year.",
            "Find road segments with high crash counts and traffic volume data.",
            "List files in the YOUR_BUCKET_NAME bucket.",
            "Search the bridge inspection PDF manuals for bridge inspection responsibility and cite the PDF name and page.",
        ],
    )

    doc.add_heading("9. Maintenance Notes", level=1)
    add_bullets(
        doc,
        [
            "If Cloud Run startup fails, check requirements.txt first for missing A2A or PDF dependencies.",
            "If Gemini Enterprise reports part_metadata errors, keep the custom request converter in agent.py.",
            "If PDF text extraction returns empty text, the PDF may be scanned images; use OCR such as Document AI or Vision OCR.",
            "If BigQuery questions fail, confirm BIGQUERY_TABLES exactly matches fully qualified table ids.",
            "If GCS PDF reading fails, confirm the Cloud Run service account has storage.objects.get permission for the bucket.",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Transportation Data A2A ADK Agent - Code Reference")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
